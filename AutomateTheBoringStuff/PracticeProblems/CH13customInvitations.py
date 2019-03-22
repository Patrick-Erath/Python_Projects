# customInvitations.py : makes a word document with a custome page for an invitation
# of each guest in the guest.txt file
import docx

doc = docx.Document('./BirthdayInvitations.docx') 
# Get guests from guests.txt
guests = [line.rstrip('\n') for line in open('guests.txt')]

# add guests to word document
# make sure font is conserved
doc.paragraphs[4].runs[1].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.paragraphs[1].text = (guests[0])

# Get the different styles
styles = []
for i in range(4):
	styles.append(doc.paragraphs[i].style)

# Get the different texts
texts = ['intro', 'name', 'address', 'date', 'time']
info = {}
for i in range(5):
	if(i!=1):
		info[texts[i]] = doc.paragraphs[i].text

n = 9
for i in range(1,len(guests)):
	# Add each part of the invitation
	doc.add_paragraph(info['intro'])
	doc.paragraphs[n-4].style = styles[0]
	doc.add_paragraph(guests[i])
	doc.paragraphs[n-3].style = styles[1]
	doc.add_paragraph(info['address'])
	doc.paragraphs[n-2].style = styles[0]
	doc.add_paragraph(info['date'])
	doc.paragraphs[n-1].style = styles[3]
	doc.add_paragraph(info['time'])
	doc.paragraphs[n].style = styles[0]

	# Add break, i.e go to next page to write next invite
	doc.paragraphs[n].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
	n+=5

doc.save('combinedInvitations.docx')